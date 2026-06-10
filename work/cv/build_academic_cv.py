from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("work/cv/Jie_Hu_Academic_CV.docx")
BLUE = RGBColor(31, 77, 120)
DARK = RGBColor(28, 39, 44)
MUTED = RGBColor(92, 103, 108)
ACCENT = RGBColor(184, 88, 66)
LIGHT = "E8EEF5"


def font(run, size=10, bold=False, italic=False, color=DARK, name="Aptos"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    return run


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def add_hyperlink(paragraph, text, url, color=ACCENT):
    part = paragraph.part
    rid = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rid)
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), str(color))
    rPr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    hyperlink.append(r)
    paragraph._p.append(hyperlink)


def section_title(doc, title):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(5)
    font(p.add_run(title.upper()), 12.5, True, color=BLUE)
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), "D7DEE4")
    border.append(bottom)
    p._p.get_or_add_pPr().append(border)


def entry(doc, title, meta, bullets=None, detail=None):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(1)
    font(p.add_run(title), 10, True)
    if meta:
        font(p.add_run("  |  " + meta), 9, italic=True, color=MUTED)
    if detail:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(2)
        p2.paragraph_format.line_spacing = 1.1
        font(p2.add_run(detail), 9.2, color=MUTED)
    for item in bullets or []:
        b = doc.add_paragraph(style="List Bullet")
        b.paragraph_format.left_indent = Inches(0.22)
        b.paragraph_format.first_line_indent = Inches(-0.12)
        b.paragraph_format.space_after = Pt(1.5)
        b.paragraph_format.line_spacing = 1.08
        font(b.add_run(item), 9.1)


def pub(doc, text, status=None, link=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.first_line_indent = Inches(-0.2)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.08
    font(p.add_run("• "), 9.2, True, color=ACCENT)
    font(p.add_run(text), 9.2)
    if status:
        font(p.add_run(f" [{status}]"), 8.8, True, color=ACCENT)
    if link:
        p.add_run("  ")
        add_hyperlink(p, "Link", link)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.62)
section.bottom_margin = Inches(0.58)
section.left_margin = Inches(0.72)
section.right_margin = Inches(0.72)
section.header_distance = Inches(0.3)
section.footer_distance = Inches(0.3)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Aptos"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
normal.font.size = Pt(9.5)
normal.font.color.rgb = DARK
normal.paragraph_format.space_after = Pt(3)
normal.paragraph_format.line_spacing = 1.12

for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
    s = styles[style_name]
    s.font.name = "Aptos"
    s._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    s._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
font(footer.add_run("Jie Hu · Academic Curriculum Vitae · June 2026"), 8, color=MUTED)

# Page 1
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(1)
font(p.add_run("JIE HU"), 25, True, color=DARK)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(5)
font(p.add_run("Researcher in Generative AI, Computer Graphics, HCI, and Digital Heritage"), 11, italic=True, color=BLUE)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(8)
font(p.add_run("Shanghai, China  |  122089111@qq.com  |  "), 9, color=MUTED)
add_hyperlink(p, "Google Scholar", "https://scholar.google.com/citations?user=OY2L2fIAAAAJ&hl=en")
font(p.add_run("  |  "), 9, color=MUTED)
add_hyperlink(p, "DOI: npj Heritage Science", "https://doi.org/10.1038/s40494-026-02649-7")

section_title(doc, "Research Profile")
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(5)
p.paragraph_format.line_spacing = 1.15
font(
    p.add_run(
        "Interdisciplinary researcher and technical artist investigating how generative AI, "
        "computer graphics, and interactive systems can support visual storytelling, creative production, "
        "and the interpretation of cultural heritage. Current work spans multimodal co-creative agents, "
        "3D Gaussian Splatting, world-model artworks, visual analytics, and AI-assisted game asset pipelines."
    ),
    9.5,
)

tbl = doc.add_table(rows=2, cols=2)
tbl.autofit = False
for row in tbl.rows:
    row.cells[0].width = Inches(1.6)
    row.cells[1].width = Inches(5.3)
    for c in row.cells:
        set_cell_margins(c)
shade(tbl.cell(0, 0), LIGHT)
shade(tbl.cell(1, 0), LIGHT)
for label, value, row in [
    ("Research Interests", "Generative AI · Human–Computer Interaction · Computer Graphics · Digital Heritage · Game Design", 0),
    ("Methods & Tools", "Multimodal LLMs · Diffusion Models · 3D Gaussian Splatting · Unity/Unreal Engine · Shader Development · Visual Analytics", 1),
]:
    font(tbl.cell(row, 0).paragraphs[0].add_run(label), 9, True, color=BLUE)
    font(tbl.cell(row, 1).paragraphs[0].add_run(value), 9)

section_title(doc, "Education")
entry(
    doc,
    "Research Exchange Student, Digital Art Laboratory",
    "Fudan University · Dec 2025–Present",
    detail="Advanced research in digital media art, interactive systems, and AI-enabled creative practice.",
)
entry(
    doc,
    "M.A. in Digital Media (Game Design)",
    "Shanghai University of Engineering Science · Sep 2023–Jun 2026",
    detail="GPA: Top 5%; First- and Second-Class Graduate Academic Scholarships. Thesis: Research on the Application of Maqiao Culture in Adventure Puzzle Games.",
)
entry(
    doc,
    "B.A. in Digital Media Art",
    "Shanghai Jian Qiao University · Sep 2019–Jun 2023",
    detail="GPA: Top 10%; consecutive university scholarships; Outstanding Graduation Thesis/Design Award.",
)

section_title(doc, "Selected Research Contributions")
entry(
    doc,
    "Digital Heritage Reconstruction and Interpretation",
    "",
    bullets=[
        "Developed UAV-photogrammetry and 3D Gaussian Splatting workflows for reconstructing historical buildings and metaverse environments.",
        "Designed immersive and participatory systems that expose uncertainty rather than presenting machine-generated heritage as unquestioned fact.",
    ],
)
entry(
    doc,
    "AI-Assisted Visual Creation",
    "",
    bullets=[
        "Built multimodal frameworks for visual narratives, comics, storyboards, and semantically guided 2D-to-3D game asset generation.",
        "Explored controllable synthesis through co-creative agents, retrieval augmentation, and neuro-symbolic systems.",
    ],
)

doc.add_page_break()

# Page 2
section_title(doc, "Publications and Manuscripts")
pub(
    doc,
    "Jie Hu, Zhaoli Jiang, and Yingfang Zhang. “Digital Regeneration of Historical Buildings Using UAV Photogrammetry and 3D Gaussian Splatting for Metaverse Environments.” npj Heritage Science, 2026.",
    "Published",
    "https://doi.org/10.1038/s40494-026-02649-7",
)
pub(
    doc,
    "Jie Hu, Jinyu Li, Zixia Wang, Zhixian Li, Kachun Chan, Zhaoli Jiang, and Yingfang Zhang. “From Pixel to Mesh: Accelerating Game Asset Creation via a Semantically-Guided 2D-to-3D Generative Pipeline.” AHFE 2026.",
    "Accepted",
)
pub(
    doc,
    "Jie Hu, Yingfang Zhang, and Zhaoli Jiang. “Innovative Application of AIGC Technology in Digital Game Design.” ICDI 2024.",
    "Published",
)
pub(
    doc,
    "Jie Hu, Zhaoli Jiang, and Yingfang Zhang. “Parametric Design Mechanism of Cultural Products Driven by AIGC.” Shoes Technology and Design, 2026.",
    "Published",
)
pub(
    doc,
    "Zhaoli Jiang, Xiaoxian Ye, and Jie Hu. “Research on the Application of Generative AI in the Digital Design of Miao Totem Symbols.” Package & Design, 2025.",
    "Published",
)
pub(
    doc,
    "Jie Hu, Zhaoli Jiang, and Yingfang Zhang. “Unweaving the Machine Scroll: A Participatory World-Model Artwork for Inspecting Historical Uncertainty in Qingming Shanghe Tu.” SIGGRAPH Asia 2026 Art Papers.",
    "Under Review",
)
pub(
    doc,
    "Jie Hu, Zhaoli Jiang, and Yingfang Zhang. “DualDraft: Enabling Zero-Experience Users to Create Consistent Visual Narratives via a Dual-Mode Co-Creative Agent.” UIST 2026.",
    "Under Review",
)
pub(
    doc,
    "Jie Hu, Zhaoli Jiang, and Yingfang Zhang. “Palimpsest of Maqiao: An Immersive VR Visual Analytics Framework through Stylized 3D Gaussian Splatting.” IEEE TVCG.",
    "Under Review",
)
pub(
    doc,
    "Jie Hu, Zhaoli Jiang, and Yingfang Zhang. “KomixFlow: A Retrieval-Augmented Multimodal Framework for Long-Horizon Visual Narrative Generation.” IEEE Transactions on Multimedia.",
    "Under Review",
)
pub(
    doc,
    "Jie Hu, Zhaoli Jiang, and Yingfang Zhang. “StoryNode: A Node-Based Visual Orchestration Framework for Automated Comic and Storyboard Generation.” International Journal of Human–Computer Interaction.",
    "Under Review",
)
pub(
    doc,
    "Jie Hu, Zhaoli Jiang, and Yingfang Zhang. “A Neuro-Symbolic Expert System for Controllable LLM-Driven Multimodal Content Synthesis.” Expert Systems with Applications.",
    "Under Review",
)
pub(
    doc,
    "Jie Hu, Zhaoli Jiang, and Yingfang Zhang. “Exploration of Revitalizing ‘Cold Heritage’ Empowered by Generative AI: A Case Study of Maqiao Pottery Particle System.” Journal of Graphics.",
    "Under Review",
)

section_title(doc, "Academic Service")
entry(doc, "Reviewer", "ACM UIST 2026", detail="Reviewed submissions for the ACM Symposium on User Interface Software and Technology.")
entry(doc, "Reviewer", "ACM ICMR 2026", detail="Reviewed submissions for the ACM International Conference on Multimedia Retrieval.")

section_title(doc, "Selected Honors and Awards")
entry(doc, "National Third Prize", "11th Future Designer National College Digital Art & Design Awards · 2023")
entry(doc, "Second Prize and Two Third Prizes", "9th Huichuang Youth Shanghai College Students Cultural and Creative Works Exhibition · 2024")
entry(doc, "Regional First Prizes", "10th and 11th Future Designer NCDA, Shanghai · 2022–2023")
entry(doc, "Outstanding Graduation Thesis/Design Award", "Shanghai Jian Qiao University · 2023")
entry(doc, "Graduate Academic Scholarships", "First- and Second-Class Awards · 2023–2024")

doc.add_page_break()

# Page 3
section_title(doc, "Research-Relevant Professional Experience")
entry(
    doc,
    "Technical Artist",
    "NetEase Games · Dec 2025–Present",
    bullets=[
        "Develop AI agents and multimodal video workflows for game production and provide rendering support for high-fidelity FPS projects.",
        "Optimize realistic materials, shaders, and scene-rendering workflows by integrating generative AI with conventional graphics pipelines.",
    ],
)
entry(
    doc,
    "Art Project Manager, Infinity Nikki",
    "Papergames · Sep 2025–Dec 2025",
    bullets=[
        "Managed cross-platform VFX production for a large-scale open-world title across PC, mobile, and PS5.",
        "Coordinated art, engineering, design, and external production teams to improve asset delivery and pipeline reliability.",
    ],
)
entry(
    doc,
    "R&D Project Manager and Technical Artist, DreamStar",
    "Tencent TiMi Studio · Mar 2025–Aug 2025",
    bullets=[
        "Established client-module testing standards and defect-analysis workflows for open-world gameplay, rendering, physics, and AI.",
        "Collaborated with engineering and art teams to resolve rendering anomalies, terrain-material issues, and development risks.",
    ],
)
entry(
    doc,
    "Marketing and Publishing Project Manager",
    "Lilith Games · Oct 2024–Mar 2025",
    bullets=[
        "Led China-region creative production and coordinated development, publishing, and external teams for campaign delivery.",
        "Used performance data and market feedback to support publishing decisions and resource allocation.",
    ],
)

section_title(doc, "Selected Projects")
entry(
    doc,
    "Unweaving the Machine Scroll",
    "Participatory World-Model Artwork · 2026",
    detail="Transforms Qingming Shanghe Tu into an inspectable machine-generated environment, enabling visitors to distinguish evidence, inferred continuity, and algorithmic speculation.",
)
entry(
    doc,
    "Maqiao Cultural Journey",
    "Master’s Thesis and Interactive Game · 2026",
    detail="A 3D adventure-puzzle game that integrates archaeological and cultural elements to support the digital revitalization of Maqiao heritage.",
)
entry(
    doc,
    "Lost Ruins – Akrosa",
    "3D Environment Design · 2023",
    detail="End-to-end Unreal Engine 5 environment production using a complete PBR workflow; received the Outstanding Graduation Thesis/Design Award.",
)

section_title(doc, "Technical Skills")
tbl = doc.add_table(rows=5, cols=2)
tbl.autofit = False
skills = [
    ("AI and Research", "Multimodal LLMs, diffusion models, LoRA, ControlNet, ComfyUI, co-creative agents, retrieval-augmented generation"),
    ("Graphics and Engines", "Unreal Engine 5, Unity, 3D Gaussian Splatting, PBR workflows, shader development, procedural modeling"),
    ("Programming", "Python, C#, C++ fundamentals, GLSL/HLSL"),
    ("Design and DCC", "Blender, Maya, 3ds Max, Substance Painter, Adobe Creative Suite"),
    ("Languages", "Chinese (native), English (research and professional working proficiency), Korean (TOPIK Level 2)"),
]
for i, (label, value) in enumerate(skills):
    tbl.rows[i].cells[0].width = Inches(1.55)
    tbl.rows[i].cells[1].width = Inches(5.35)
    for c in tbl.rows[i].cells:
        set_cell_margins(c)
    shade(tbl.cell(i, 0), LIGHT)
    font(tbl.cell(i, 0).paragraphs[0].add_run(label), 8.8, True, color=BLUE)
    font(tbl.cell(i, 1).paragraphs[0].add_run(value), 8.8)

section_title(doc, "Links")
p = doc.add_paragraph()
font(p.add_run("Google Scholar: "), 9, True)
add_hyperlink(p, "scholar.google.com/citations?user=OY2L2fIAAAAJ", "https://scholar.google.com/citations?user=OY2L2fIAAAAJ&hl=en")
p.add_run("   ")
font(p.add_run("Email: "), 9, True)
font(p.add_run("122089111@qq.com"), 9)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
